"""
Exercise 45: You Make a Game
    You need to start learning to feed yourself. Hopefully as you have
    worked through this book, you have learned that all the information
    you need is on the internet. You just have to go search for it. The
    only thing you have been missing are the right words and what to
    look for when you search. Now you should have a sense of it, so
    it's about time you struggle through a big project and tried to get
    it working.
"""

"""
Code Style
    * Give your code vertical space so people can read it. You will find
    some very bad programmers who are able to write reasonable code but
    who do not add any spaces. This si bad style in any language
    because the human eye and brain use space and vertical alignment to
    scan and separate visual elements. Not having space is the same as
    giving your code an awesome camouflage paint job.

    * If you can't read it out loud, it's probably hard to read. If you
    are having a problem making something easy to use, try reading it
    out loud. Not only does this force difficult passage and things to
    change for readability.

    * Try to do what other people are doing in Python until you find
    your own style.

    * Once you find your own style, do not be a jerk about it. Working
    with other people's code is part of being a programmer, and other
    people have really bad taste. Trust me, you will probably have
    really bad taste too and even realize it.
"""

"""
Good Comments
    * Programmers will tell you that your code should be readable
    enough that you do not need comments. They'll then tell you in
    their most official sounding voice, "Ergo one should never write
    comments or documentation. QED". Those programmers are either
    consultants who get paid more if other people can't use their code,
    or incompetents who tend to never work with other people. Ignore
    them and write comments.

    * When you write comments, describe why you are doing what you are
    doing. The code already says how, but why you did things way you
    did is more important.

    * When you write doc comments for your functions, make the comments
    documentation for someone who will have to use your code. You do
    not have to go crazy, but a nice little sentence about what someone
    can do with that function helps a lot.

    * While comments are good, too many are bad, and you have to
    maintain them. Keep your comments relatively short and to the
    point, and if you change a function, review the comment to make
    sure it's still correct.
"""
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.add_picture('monty-truth.png', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'


class records(object):
    def __init__(self, Qty=-1, Ids=-1, Desc=None):
        self.qty = Qty
        self.id = Ids
        self.desc = Desc

r1 = records(1, 101, 'Spam')
r2 = records(2, 42, 'Eggs')
r3 = records(3, 631, 'Spam, spam, eggs, and spam')
recordset = [r1, r2, r3]

for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item.qty)
    row_cells[1].text = str(item.id)
    row_cells[2].text = item.desc

document.add_page_break()

document.save('demo.docx')
